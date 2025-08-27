import os
from datetime import datetime
from logger.log import CustomLogger
from core.constants import Margins, Dimensions, FolderName
from core.snap import Reference
from docx import Document
from docx.shared import Inches, Mm
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH


class RCA:
    """
    The RCA (Root Cause Analysis) class is responsible for generating
    Word documents that contain findings from a security analysis. It
    takes a list of `Reference` objects, processes them, and creates a
    formatted document with headings, scores, and images.
    """

    def __init__(self, ref: list[Reference]):
        """
        Initializes the RCA class.

        Args:
            ref (list[Reference]): A list of references containing information
            to be documented, including headings, scores,
            and paths to screenshots.
        """
        self._logger = CustomLogger()
        self.ref = ref

    def document_findings(self):
        """
        Generates a Word document with the findings.

        This method sets up the document's page size and margins, then iterates
        through each `Reference` object to add a heading, score, and screenshot
        to the document. It includes error handling to ensure the process continues
        even if an issue occurs with a single reference.

        The final document is saved with a timestamped filename in a dedicated
        folder.
        """
        doc = Document()
        self._logger.info("Setting Doc size to A4")
        section = doc.sections[0]
        section.page_height = Mm(Dimensions.HEIGHT.value)
        section.page_width = Mm(Dimensions.WIDTH.value)
        section.top_margin = Inches(Margins.ALL.value)
        section.bottom_margin = Inches(Margins.ALL.value)
        section.left_margin = Inches(Margins.ALL.value)
        section.right_margin = Inches(Margins.ALL.value)

        for r in self.ref:
            try:
                heading = r.heading
                score = r.score
                path = r.path

                head: Paragraph = doc.add_heading(heading, 0)
                head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(score)
                doc.add_picture(path, width=Inches(Dimensions.IMG_WIDTH.value))
                doc.add_paragraph("")
            except Exception as err:
                self._logger.error(f"An unexpected error occurred: {err}")

        file_name = f"RCA_{datetime.now().strftime('%Y_%m_%d_%H_%M_%S')}.docx"
        folder_name = FolderName.DOC_FOLDER_NAME.value
        destination_folder = os.path.join(os.getcwd(), folder_name)
        os.makedirs(destination_folder, exist_ok=True)
        full_path = os.path.join(destination_folder, file_name)
        self._logger.success("File successfully saved at location: ")
        self._logger.success(f"{full_path}")
        doc.save(full_path)
